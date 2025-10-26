from docx import Document
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ContratoService:

    def ler_texto_docx(self, caminho_arquivo):
        doc = Document(caminho_arquivo)
        return "\n".join([p.text for p in doc.paragraphs])

    def extrair_blocos_principais(self, texto):
        """Identifica todos os blocos principais como {TIPO_LOCATARIA}, {TIPO_FIADOR}, etc."""
        padrao = r"\{([A-Z_]+)\}([\s\S]*?)\{/\1\}"
        return re.findall(padrao, texto)

    def processar_opcoes_bloco(self, bloco_texto, tipo_selecionado):
        padrao_opcao = r"\{\{([A-Z_]+)\}\}([\s\S]*?)\{\{/\1\}\}"
        opcoes = re.findall(padrao_opcao, bloco_texto)
        if not opcoes:
            return bloco_texto
        for nome_opcao, conteudo in opcoes:
            if nome_opcao == tipo_selecionado:
                return conteudo
        return ""

    def aplicar_json_em_estrutura(self, texto, data_fields):
        """Percorre o documento e aplica a lógica de estrutura com base nas opções do JSON."""
        blocos = self.extrair_blocos_principais(texto)
        for nome_bloco, conteudo in blocos:
            # Exemplo: nome_bloco = "TIPO_LOCATARIA"
            if nome_bloco in data_fields:
                tipo_valor = data_fields[nome_bloco]["valor"]  # ex: PESSOA_FISICA
                novo_conteudo = self.processar_opcoes_bloco(conteudo, tipo_valor)
                # Substitui o bloco original no texto
                padrao_bloco = re.compile(rf"\{{{nome_bloco}\}}([\s\S]*?)\{{/{nome_bloco}\}}", re.DOTALL)
                texto = padrao_bloco.sub(novo_conteudo, texto, count=1)
        # Substitui placeholders diretos
        for chave, valor in data_fields.items():
            if valor.get("type", "").lower() == "value":
                texto = texto.replace(f"{{{chave}}}", valor["valor"])

        texto = re.sub(r'\n{3,}', '\n\n', texto.strip())
        return texto

    def remover_placeholders_restantes(self, texto: str) -> str:
        texto = re.sub(r"\{[A-Z0-9_]+\}", "", texto)
        texto = re.sub(r"[ \t]+", " ", texto)
        return re.sub(r"\n{3,}", "\n\n", texto.strip())

    def injetar_texto_em_modelo(self, caminho_modelo, texto_final, caminho_saida):
        modelo = Document(caminho_modelo)

        while modelo.paragraphs:
            p = modelo.paragraphs[0]._element
            p.getparent().remove(p)

        for linha in texto_final.split("\n"):
            p = modelo.add_paragraph(linha.strip())
            p.style = modelo.styles["Normal"]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        modelo.save(caminho_saida)
        return caminho_saida

    def processar_contrato(self, data_fields: dict) -> str:
        caminho_modelo = "contrato/ID_1.docx"
        caminho_saida = "result/contrato_gerado.docx"
        texto_original = self.ler_texto_docx(caminho_modelo)
        texto_processado = self.aplicar_json_em_estrutura(texto_original, data_fields)
        texto_limpo = self.remover_placeholders_restantes(texto_processado)
        return self.injetar_texto_em_modelo(caminho_modelo, texto_limpo, caminho_saida)
